from fastapi import FastAPI, Request
from fastapi.middleware.cors import CORSMiddleware
import requests
from pypdf import PdfReader
import base64
import io
from docx import Document
from backend.database import init_db
import sqlite3
from backend.database import DB_PATH
import hashlib
from datetime import datetime

app = FastAPI()

origins = [
    "http://localhost",
    "http://localhost:5173",
    "http://172.16.16.106:5173"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

init_db()

@app.get("/")
async def read_root(request: Request):
    ip = request.client.host
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("INSERT INTO visits (ip, dateTime) VALUES (?,?)", (ip, timestamp))
        conn.commit() 
        status = "success"
        message = "Welcome"

    except sqlite3.Error as e:
        status = "error"
        message = e
    finally:
        conn.close()
    
    return {"status": status, "message": message}

@app.get("/models")
def models():
    url = "http://172.16.16.106:1234/v1/models"
    response = requests.get(url)
    return response.json()

@app.post("/responses")
async def responses(payload: dict):
    url = "http://172.16.16.106:1234/v1/responses"
    headers = {"Content-Type": "application/json"}
    selected_model = payload.get("selectedModel")
    input_text = payload.get("input")
    temperature = payload.get("temperature")
    previous_response = payload.get("previousResponse")
    file = payload.get("file")
    image = payload.get("image")
    timestamp = datetime.now().strftime("%Y-%m-%d")

    if file:
        extracted_text = ""
        docx = file.startswith("data:application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        if "," in file:
            file = file.split(",")[1]
            
        file_bytes = base64.b64decode(file)
        file_stream = io.BytesIO(file_bytes)
        
        if docx:
            doc = Document(file_stream)
            for para in doc.paragraphs:
                if para.text:
                    extracted_text += para.text + "\n"
                
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        extracted_text += "\n" + "\t".join(row_text) + "\n"

            max_chars = 11000
            if len(extracted_text) > max_chars:
                extracted_text = extracted_text[:max_chars]
                
        else:
            reader = PdfReader(file_stream)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    extracted_text += text + "\n"

        max_chars = 12000
        if len(extracted_text) > max_chars:
            extracted_text = extracted_text[:max_chars]
        
        data = {
            "model": selected_model,
            "input": 
            [
                {
                    "role": "user",
                    "content": 
                    [
                        { 
                            "type": "input_text", 
                            "text": input_text 
                        },
                        { 
                            "type": "input_text",
                            "text": extracted_text
                        }
                    ]
                }
            ]
        }
    
    elif image:
        data = {
            "model": selected_model,
            "input": 
            [
                {
                    "role": "user",
                    "content": 
                    [
                        { 
                            "type": "input_text", 
                            "text": input_text 
                        },
                        { 
                            "type": "input_image",
                            "image_url": image
                        }
                    ]
                }
            ]
        }

    else:
        data = {
            "model": selected_model,
            "input": 
            [
                {
                    "role": "user",
                    "content": 
                    [
                        { 
                            "type": "input_text", 
                            "text": input_text 
                        }
                    ]
                }
            ],
            "temperature": temperature
        }

    if previous_response:
        data["previous_response_id"] = previous_response
        
    response = requests.post(url, headers=headers, json=data)
    responseData = response.json()

    tokens = responseData.get("usage", {}).get("total_tokens", 0)
    
    try:
            conn = sqlite3.connect(DB_PATH)
            cursor = conn.cursor()

            cursor.execute("INSERT INTO tokens (dateTime, amount) VALUES (?, ?)",(timestamp, tokens))
            conn.commit()
    except sqlite3.Error as e:
            print(f"Datenbankfehler beim Token-Log: {e}")
    finally:
        conn.close()
   
    return response.json()
    
@app.post("/register")
def register(data: dict):
    username = data.get("username")
    password = data.get("password")

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try: 
        cursor.execute("SELECT * FROM users WHERE username = ?", (username,))
        existing_user = cursor.fetchone()

        if existing_user:
            status = "error"
            message = "Benutzername existiert bereits"
        
        else:
            password_bytes = password.encode('utf-8')
            hashed_password = hashlib.sha256(password_bytes).hexdigest()

            cursor.execute("INSERT INTO users (username, password) VALUES (?, ?)", (username, hashed_password))
            conn.commit()
            status = "success"
            message = f"Registrierung erfolgreich! Willkommen, {username}."
    except sqlite3.Error as e:
        status = "error"
        message = f"Datenbankfehler: {str(e)}"
    finally:
        conn.close()
        
    return {"status": status, "message": message}

@app.post("/login")
def login(data: dict):
    username = data.get("username")
    password = data.get("password")
    hashed_password = hashlib.sha256(password.encode('utf-8')).hexdigest()

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    isAdmin = False
    try:
        cursor.execute("SELECT * FROM users WHERE username = ? AND password = ?", (username, hashed_password))
        user = cursor.fetchone()
        if user:
            status = "success"
            message = f"Willkommen zurück, {username}!"
            if user[3] == "admin":
                isAdmin = True
        else:  
            status = "error"
            message = "Falscher Benutzername oder Passwort!"
    except sqlite3.Error as e:
        status = "error"
        message = f"Datenbankfehler: {str(e)}"
    finally:
        conn.close()
    
    return {"status": status, "message": message, "isAdmin": isAdmin}
   
@app.post("/report")
async def report(data: dict):
    username = data.get("username")
    report_text = data.get("reportText")
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("PRAGMA foreign_keys = ON;")

        cursor.execute("SELECT id FROM users WHERE username = ?", (username,))
        row = cursor.fetchone()
        if row:
            userId = row[0]
            cursor.execute("INSERT INTO reports (reportText, createdAt, userId) VALUES (?,?,?)", (report_text, timestamp, userId))
            conn.commit() 
            status = "success"
            message = "Bericht eingetragen"
        else:
            status = "error"
            message = "Benutzer fehlt"

    except sqlite3.Error as e:
        status = "error"
        message = str(e)
    finally:
        conn.close()
    
    return {"status": status, "message": message}

@app.post("/saveTokens")
async def tokens(data: dict):
    tokens = data.get("tokens")
    timestamp = datetime.now().strftime("%Y-%m-%d")

    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()

        cursor.execute("INSERT INTO tokens (dateTime, amount) VALUES (?,?)", (timestamp, tokens))
        conn.commit() 
        status = "success"
        message = "Tokens gespeichert"

    except sqlite3.Error as e:
        status = "error"
        message = str(e)
    finally:
        conn.close()
    
    return {"status": status, "message": message}

@app.get("/infos")
async def infos():
    timestamp = datetime.now().strftime("%Y-%m-%d")

    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()

        cursor.execute("SELECT COUNT(*) FROM tokens WHERE date(dateTime) = ?", (timestamp,))
        resultPromptsDaily = cursor.fetchone()
        promptsDaily = resultPromptsDaily[0]

        cursor.execute("SELECT COUNT(*) FROM tokens")
        resultPromptsAll = cursor.fetchone()
        promptsAll = resultPromptsAll[0]

        cursor.execute("SELECT SUM(amount) FROM tokens WHERE date(dateTime) = ?", (timestamp,))
        resultDaily = cursor.fetchone()
        tokensDaily = resultDaily[0]

        cursor.execute("SELECT SUM(amount) FROM tokens")
        resultAll = cursor.fetchone()
        tokensAll = resultAll[0]

        cursor.execute("SELECT MIN(dateTime) FROM tokens")
        resultOldestDate = cursor.fetchone()
        oldestDate = resultOldestDate[0]

        cursor.execute("SELECT COUNT(*) FROM users")
        resultUsersCount = cursor.fetchone()
        usersCount = resultUsersCount[0]    
        status = "success"
        message = "Get Tokens"

    except sqlite3.Error as e:
        status = "error"
        message = str(e)
    finally:
        conn.close()
    
    return {"status": status, "message": message, "promptsDaily": promptsDaily, "promptsAll": promptsAll, "tokensDaily": tokensDaily, "tokensAll": tokensAll, "oldestDate": oldestDate, "usersCount": usersCount}